import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_pleading_table(document):
    """
    Creates and formats a 28-line, 3-column table for pleading paper.
    Returns the created table object.
    """
    # Table with 3 columns: [Numbers], [Plaintiff Info], [Case Info]
    table = document.add_table(rows=28, cols=3)
    
    # Set column widths for 8.5" paper with 0.5" margins on each side
    table.columns[0].width = Inches(0.5)
    table.columns[1].width = Inches(3.5)
    table.columns[2].width = Inches(3.5)

    # Format borders: vertical lines between each column
    tbl = table._tbl
    tblPr = tbl.tblPr
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ["top", "left", "bottom", "right", "insideH"]:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tbl_borders.append(border)
    
    inside_v_border = OxmlElement('w:insideV')
    inside_v_border.set(qn('w:val'), 'single')
    inside_v_border.set(qn('w:sz'), '4')
    tbl_borders.append(inside_v_border)
    tblPr.append(tbl_borders)

    # Populate line numbers and set default row formatting
    for i in range(28):
        table.rows[i].height = Pt(24) # Enforce row height for double spacing
        p_num = table.cell(i, 0).paragraphs[0]
        p_num.add_run(str(i + 1))
        
        # Set default text columns to be double-spaced
        for j in range(1, 3):
            p_text = table.cell(i, j).paragraphs[0]
            p_text.paragraph_format.line_spacing = Pt(24)
            p_text.paragraph_format.space_after = Pt(0)
    
    return table

def create_final_pleading_template(filename="pleading_template_final.docx"):
    """
    Creates a clean, professional Word document formatted for CA pleading paper.
    """
    document = docx.Document()
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # --- Create and Populate Page 1 ---
    table_page1 = create_pleading_table(document)

    # Helper to set single spacing for header/caption text
    def set_single_space(cell):
        cell.paragraphs[0].paragraph_format.line_spacing = Pt(12)

    # Attorney Info
    attorney_lines = [
        "[Name or Firm Name], SBN [State Bar #]",
        "[Street Address]",
        "[City, State, Zip Code]",
        "Telephone: [Telephone #]",
        "Attorney for: [Party Name]"
    ]
    for i, line in enumerate(attorney_lines):
        cell = table_page1.cell(i, 1)
        set_single_space(cell)
        cell.text = line

    # Court Title (Centered)
    court_cell = table_page1.cell(7, 1)
    court_cell.merge(table_page1.cell(7, 2))
    set_single_space(court_cell)
    court_cell.text = "SUPERIOR COURT OF THE STATE OF CALIFORNIA\nCOUNTY OF [COUNTY NAME]"
    court_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    court_cell.paragraphs[0].runs[0].bold = True
    
    # --- Caption Area ---
    # Column 1: Plaintiff Info
    p_plaintiff_name = table_page1.cell(10, 1)
    set_single_space(p_plaintiff_name)
    p_plaintiff_name.text = "THE PEOPLE OF THE STATE OF CALIFORNIA,"

    p_plaintiff_label = table_page1.cell(12, 1)
    set_single_space(p_plaintiff_label)
    p_plaintiff_label.text = "          Plaintiff,"

    p_vs = table_page1.cell(14, 1)
    set_single_space(p_vs)
    p_vs.text = "vs."
    p_vs.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_defendant_name = table_page1.cell(16, 1)
    set_single_space(p_defendant_name)
    p_defendant_name.text = "[DEFENDANT'S FULL NAME],"
    p_defendant_name.paragraphs[0].runs[0].all_caps = True
    
    p_defendant_label = table_page1.cell(18, 1)
    set_single_space(p_defendant_label)
    p_defendant_label.text = "          Defendant."

    # Column 2: Case Info
    p_case_num = table_page1.cell(10, 2)
    set_single_space(p_case_num)
    p_case_num.text = "Case No.: [CASE NUMBER]"

    p_doc_title = table_page1.cell(12, 2)
    set_single_space(p_doc_title)
    run = p_doc_title.paragraphs[0].add_run("[TITLE OF DOCUMENT]")
    run.bold = True
    run.all_caps = True

    # --- Create Page 2 ---
    document.add_page_break()
    create_pleading_table(document)
    
    try:
        document.save(filename)
        print(f"Successfully created '{filename}'")
    except Exception as e:
        print(f"Error saving file: {e}")

# Run the function to generate your template
create_final_pleading_template()