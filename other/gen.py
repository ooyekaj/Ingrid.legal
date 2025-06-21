import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_pleading_paper(filename="pleading_paper_template.docx", document_title="PLEADING PAPER"):
    """
    Creates a blank Word document with 28-line pleading paper formatting
    that meets California court requirements.

    Args:
        filename (str): The name of the Word file to create.
        document_title (str): The title to appear in the footer.
    """
    # 1. Create a new blank document
    document = docx.Document()

    # 2. Set document font to Times New Roman 12pt
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # 3. Set page margins for standard pleading paper (1 inch all around)
    sections = document.sections
    for section in sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # 4. Create a table with 28 rows and 2 columns
    num_lines = 28
    table = document.add_table(rows=num_lines, cols=2)

    # 5. Set the column widths (6.5" total width for 8.5" paper with 1" margins)
    table.columns[0].width = Inches(0.5)  # Line numbers column
    table.columns[1].width = Inches(6.0)  # Text column

    # 6. Populate the table with line numbers and format them
    for i in range(num_lines):
        # --- Number Column ---
        cell_num = table.cell(i, 0)
        p_num = cell_num.paragraphs[0]
        p_num.add_run(str(i + 1))
        p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_num.paragraph_format.space_after = Pt(0)
        p_num.paragraph_format.line_spacing = Pt(24)  # Double spacing (24pt for 12pt font)
        
        # --- Text Column ---
        cell_text = table.cell(i, 1)
        p_text = cell_text.paragraphs[0]
        p_text.paragraph_format.space_after = Pt(0)
        p_text.paragraph_format.line_spacing = Pt(24)  # Double spacing (24pt for 12pt font)

    # 7. Customize table borders to create the single vertical line
    tbl = table._tbl
    tblPr = tbl.tblPr 

    # Create a tcBorders element
    tbl_borders = OxmlElement('w:tblBorders')

    # Define each border type and set its value to "nil" to hide it
    for border_name in ["top", "left", "bottom", "right", "insideH"]:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tbl_borders.append(border)

    # Define the inside vertical border to be a single solid line
    inside_v_border = OxmlElement('w:insideV')
    inside_v_border.set(qn('w:val'), 'single')
    inside_v_border.set(qn('w:sz'), '4')
    inside_v_border.set(qn('w:space'), '0')
    inside_v_border.set(qn('w:color'), 'auto')
    tbl_borders.append(inside_v_border)

    # Append the border definitions to the table properties
    tblPr.append(tbl_borders)

    # 8. Add footer with document title and separator line
    section = document.sections[0]
    footer = section.footer
    
    # Clear any existing content
    footer.paragraphs[0].clear()
    
    # Add separator line
    separator = footer.paragraphs[0]
    separator.add_run("_" * 80)  # Create a line of underscores
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add document title
    title_para = footer.add_paragraph()
    title_run = title_para.add_run(document_title)
    title_run.font.size = Pt(10)  # At least 10-point font
    title_run.font.name = 'Times New Roman'
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 9. Save the document
    try:
        document.save(filename)
        print(f"Successfully created '{filename}' with California pleading paper formatting")
        print(f"Features:")
        print(f"- 8.5 x 11 inch paper with 1-inch margins")
        print(f"- Times New Roman 12pt font")
        print(f"- 28 numbered lines with double spacing")
        print(f"- Footer with document title: '{document_title}'")
    except Exception as e:
        print(f"Error saving file: {e}")

# Run the function to create the document
if __name__ == "__main__":
    create_pleading_paper()