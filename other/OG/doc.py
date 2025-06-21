from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_word_doc_with_lines(output_filename):
    """
    Creates a Word document with two vertical lines, each one inch from the page edge,
    using a corrected and compatible method.
    """
    document = Document()
    section = document.sections[0]

    # Set Page Size and Margins
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Work in the Header to ensure the lines are visible
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    def add_vertical_line_in_header(paragraph, horiz_pos_inch, page_height_inch):
        """
        Adds a floating vertical line using basic, compatible XML objects.
        This corrects the qn() function call error.
        """
        # EMU is the internal unit Word uses. 914400 EMUs per inch.
        horiz_pos_emu = str(int(horiz_pos_inch * 914400))
        height_emu = str(int(page_height_inch * 914400))

        # Create the XML structure for a drawing object (anchor)
        anchor = OxmlElement('wp:anchor')
        # --- CORRECTED: All qn() calls now have the required 'wp:' prefix ---
        anchor.set(qn('wp:distT'), "0")
        anchor.set(qn('wp:distB'), "0")
        anchor.set(qn('wp:distL'), "0")
        anchor.set(qn('wp:distR'), "0")
        anchor.set(qn('wp:simplePos'), "0")
        anchor.set(qn('wp:relativeHeight'), "23169520")
        anchor.set(qn('wp:behindDoc'), "0")
        anchor.set(qn('wp:locked'), "0")
        anchor.set(qn('wp:layoutInCell'), "1")
        anchor.set(qn('wp:allowOverlap'), "1")

        # Set horizontal position relative to the page
        simplePos = OxmlElement('wp:simplePos')
        simplePos.set('x', '0')
        simplePos.set('y', '0')
        positionH = OxmlElement('wp:positionH')
        positionH.set(qn('wp:relativeTo'), 'page')
        posOffsetH = OxmlElement('wp:posOffset')
        posOffsetH.text = horiz_pos_emu
        positionH.append(posOffsetH)

        # Set vertical position relative to the page
        positionV = OxmlElement('wp:positionV')
        positionV.set(qn('wp:relativeTo'), 'page')
        posOffsetV = OxmlElement('wp:posOffset')
        posOffsetV.text = "0" # Top of the page
        positionV.append(posOffsetV)

        # Define the size of the drawing object
        extent = OxmlElement('wp:extent')
        extent.set('cx', "0")
        extent.set('cy', height_emu)
        docPr = OxmlElement('wp:docPr')
        docPr.set('id', '1')
        docPr.set('name', 'Vertical Line')

        # Create the actual graphic shape
        graphic = OxmlElement('a:graphic')
        graphicData = OxmlElement('a:graphicData')
        graphicData.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/main')
        
        # Define the shape as a line
        sp = OxmlElement('a:sp')
        spPr = OxmlElement('a:spPr')
        prstGeom = OxmlElement('a:prstGeom')
        prstGeom.set('prst', 'line')
        avLst = OxmlElement('a:avLst')
        prstGeom.append(avLst)
        
        # Define the line's appearance (1pt, black)
        ln = OxmlElement('a:ln')
        ln.set('w', '12700')
        solidFill = OxmlElement('a:solidFill')
        srgbClr = OxmlElement('a:srgbClr')
        srgbClr.set('val', '000000')
        solidFill.append(srgbClr)
        ln.append(solidFill)

        # Assemble the shape
        spPr.append(prstGeom)
        spPr.append(ln)
        sp.append(spPr)
        graphicData.append(sp)
        graphic.append(graphicData)

        # Assemble the final anchor object
        anchor.append(simplePos)
        anchor.append(positionH)
        anchor.append(positionV)
        anchor.append(extent)
        anchor.append(docPr)
        anchor.append(graphic)
        
        # Add the complete drawing object to the paragraph
        paragraph._p.append(anchor)

    # --- Add the two lines to the document header ---
    page_height = 11
    page_width = 8.5

    # Add the left line 1 inch from the page edge
    add_vertical_line_in_header(p, 1.0, page_height)
    
    # Add the right line 1 inch from the right page edge
    add_vertical_line_in_header(p, 7.5, page_height)
    
    document.save(output_filename)


if __name__ == '__main__':
    output_docx_name = "word_document_with_lines.docx"
    create_word_doc_with_lines(output_docx_name)
    print(f"Successfully created '{output_docx_name}' with vertical lines 1 inch from each side.")