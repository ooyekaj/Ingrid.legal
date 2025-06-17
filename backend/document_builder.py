from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List, Optional, Any
import io

class LegalDocumentBuilder:
    """
    A builder class for creating legal document shells with proper formatting
    based on court rules and requirements.
    """
    
    def __init__(self):
        self.case_info = {}
        self.formatting = {
            'font_family': 'Times New Roman',
            'font_size': 12,
            'line_spacing': 2.0,  # Double spacing
            'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.5, 'right': 1.0}
        }
        self.sections = []
        self.page_settings = {'pages': 1}
        self.document_type = ""
        self.court_rules = []
        
    def set_case_info(self, plaintiff: str, defendant: str, case_number: str, court: str, 
                     document_type: str = "", judge: str = ""):
        """Configure case details for the document header"""
        self.case_info = {
            'plaintiff': plaintiff,
            'defendant': defendant,
            'case_number': case_number,
            'court': court,
            'document_type': document_type,
            'judge': judge
        }
        self.document_type = document_type
        return self
        
    def set_formatting(self, font_family: Optional[str] = None, font_size: Optional[int] = None, 
                      line_spacing: Optional[float] = None, margins: Optional[Dict[str, float]] = None):
        """Configure fonts, spacing, and margins"""
        if font_family:
            self.formatting['font_family'] = font_family
        if font_size:
            self.formatting['font_size'] = font_size
        if line_spacing:
            self.formatting['line_spacing'] = line_spacing
        if margins:
            self.formatting['margins'].update(margins)
        return self
        
    def add_section(self, title: str, lines_count: int = 28, include_intro_note: bool = False,
                   content_template: Optional[str] = None):
        """Add configurable sections to the document"""
        section = {
            'title': title,
            'lines_count': lines_count,
            'include_intro_note': include_intro_note,
            'content_template': content_template
        }
        self.sections.append(section)
        return self
        
    def set_page_count(self, pages: int):
        """Control document length"""
        self.page_settings['pages'] = pages
        return self
        
    def add_court_rules(self, rules: List[Dict[str, Any]]):
        """Add court rules that affect document formatting"""
        self.court_rules = rules
        return self
        
    def _apply_formatting_rules(self):
        """Parse court rules and apply formatting requirements"""
        for rule in self.court_rules:
            rule_text = rule.get('text', '').lower()
            rule_name = rule.get('name', '').lower()
            
            # Parse formatting requirements from rules
            if 'two-column' in rule_text or 'separate statement' in rule_name:
                # Special handling for separate statement format
                self._configure_separate_statement_format()
            elif 'page limit' in rule_text or 'may not exceed' in rule_text:
                # Extract page limits
                import re
                page_match = re.search(r'(\d+)\s*page', rule_text)
                if page_match:
                    max_pages = int(page_match.group(1))
                    if self.page_settings['pages'] > max_pages:
                        self.page_settings['pages'] = max_pages
            elif 'font' in rule_text:
                # Font requirements
                if 'times new roman' in rule_text:
                    self.formatting['font_family'] = 'Times New Roman'
                elif 'arial' in rule_text:
                    self.formatting['font_family'] = 'Arial'
            elif 'double spac' in rule_text:
                self.formatting['line_spacing'] = 2.0
            elif 'single spac' in rule_text:
                self.formatting['line_spacing'] = 1.0
                
    def _configure_separate_statement_format(self):
        """Configure for separate statement two-column format"""
        # Add specific section for separate statement if it's that document type
        if 'separate statement' in self.document_type.lower():
            self.add_section(
                "SEPARATE STATEMENT OF UNDISPUTED MATERIAL FACTS",
                lines_count=50,
                content_template="two_column_facts"
            )
            
    def _create_document_header(self, doc: Document):
        """Create the standard legal document header"""
        # Set up header with case information
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Case caption
        if self.case_info:
            run = header_para.add_run(f"{self.case_info.get('plaintiff', '[PLAINTIFF NAME]')}")
            run.font.name = self.formatting['font_family']
            run.font.size = Pt(self.formatting['font_size'])
            run.bold = True
            
            header_para.add_run('\n\nvs.\n\n')
            
            run = header_para.add_run(f"{self.case_info.get('defendant', '[DEFENDANT NAME]')}")
            run.font.name = self.formatting['font_family']
            run.font.size = Pt(self.formatting['font_size'])
            run.bold = True
            
            # Case number and court
            info_para = doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            info_text = f"Case No. {self.case_info.get('case_number', '[CASE NUMBER]')}\n"
            info_text += f"{self.case_info.get('court', '[COURT NAME]')}"
            if self.case_info.get('judge'):
                info_text += f"\nAssigned to Hon. {self.case_info.get('judge')}"
            
            run = info_para.add_run(info_text)
            run.font.name = self.formatting['font_family']
            run.font.size = Pt(self.formatting['font_size'])
            
            # Document title
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.add_run('\n')
            run = title_para.add_run(self.case_info.get('document_type', '[DOCUMENT TYPE]').upper())
            run.font.name = self.formatting['font_family']
            run.font.size = Pt(self.formatting['font_size'])
            run.bold = True
            run.underline = True
            
    def _create_two_column_table(self, doc: Document, section: Dict[str, Any]):
        """Create a two-column table for separate statements"""
        # Add section title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(section['title'])
        run.font.name = self.formatting['font_family']
        run.font.size = Pt(self.formatting['font_size'])
        run.bold = True
        
        doc.add_paragraph()  # Space
        
        # Create table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Set column widths
        for cell in table.columns[0].cells:
            cell.width = Inches(3.0)
        for cell in table.columns[1].cells:
            cell.width = Inches(3.0)
            
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "UNDISPUTED MATERIAL FACT"
        header_cells[1].text = "EVIDENCE SUPPORTING FACT"
        
        # Format header
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = self.formatting['font_family']
                    run.font.size = Pt(self.formatting['font_size'])
                    run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add sample rows
        for i in range(1, 11):  # Add 10 sample fact rows
            row_cells = table.add_row().cells
            row_cells[0].text = f"{i}. [Insert undisputed material fact here]"
            row_cells[1].text = f"[Insert evidence citation here - e.g., Declaration of [Name], ¶ {i}]"
            
            # Format cells
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = self.formatting['font_family']
                        run.font.size = Pt(self.formatting['font_size'])
                        
    def _add_standard_section(self, doc: Document, section: Dict[str, Any]):
        """Add a standard section with placeholder content"""
        # Section heading
        heading_para = doc.add_paragraph()
        heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading_para.add_run(section['title'])
        run.font.name = self.formatting['font_family']
        run.font.size = Pt(self.formatting['font_size'])
        run.bold = True
        
        doc.add_paragraph()  # Space
        
        # Add placeholder content lines
        lines_added = 0
        target_lines = section['lines_count']
        
        if section.get('include_intro_note'):
            intro_para = doc.add_paragraph()
            intro_text = f"[This section should contain the {section['title'].lower()} for your {self.document_type}. Please replace this placeholder text with your substantive content.]"
            run = intro_para.add_run(intro_text)
            run.font.name = self.formatting['font_family']
            run.font.size = Pt(self.formatting['font_size'])
            run.italic = True
            lines_added += 3
            
        # Add placeholder lines
        while lines_added < target_lines:
            para = doc.add_paragraph()
            run = para.add_run("[Insert content here]")
            run.font.name = self.formatting['font_family']
            run.font.size = Pt(self.formatting['font_size'])
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = self.formatting['line_spacing']
            lines_added += 1
            
    def render(self) -> Document:
        """
        Create the actual DOCX document using all configured settings
        Returns a python-docx Document object
        """
        # Apply formatting rules from court requirements
        self._apply_formatting_rules()
        
        # Create new document
        doc = Document()
        
        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(self.formatting['margins']['top'])
            section.bottom_margin = Inches(self.formatting['margins']['bottom'])
            section.left_margin = Inches(self.formatting['margins']['left'])
            section.right_margin = Inches(self.formatting['margins']['right'])
            
        # Create document header
        self._create_document_header(doc)
        
        # Add sections
        for section in self.sections:
            doc.add_page_break()
            
            if section.get('content_template') == 'two_column_facts':
                self._create_two_column_table(doc, section)
            else:
                self._add_standard_section(doc, section)
                
        # Set default font for entire document
        style = doc.styles['Normal']
        font = style.font
        font.name = self.formatting['font_family']
        font.size = Pt(self.formatting['font_size'])
        
        return doc
        


class MotionForSummaryJudgmentBuilder(LegalDocumentBuilder):
    """Specialized builder for Motion for Summary Judgment documents"""
    
    def __init__(self):
        super().__init__()
        self.document_type = "Motion for Summary Judgment"
        
    def configure_standard_sections(self):
        """Add standard sections for summary judgment motion"""
        self.add_section("INTRODUCTION", lines_count=15, include_intro_note=True)
        self.add_section("STATEMENT OF FACTS", lines_count=25, include_intro_note=True)
        self.add_section("ARGUMENT", lines_count=40, include_intro_note=True)
        self.add_section("CONCLUSION", lines_count=8, include_intro_note=True)
        return self


class SeparateStatementBuilder(LegalDocumentBuilder):
    """Specialized builder for Separate Statement documents"""
    
    def __init__(self):
        super().__init__()
        self.document_type = "Separate Statement of Undisputed Material Facts"
        
    def configure_two_column_format(self):
        """Configure for two-column separate statement format"""
        self.add_section(
            "SEPARATE STATEMENT OF UNDISPUTED MATERIAL FACTS",
            lines_count=50,
            content_template="two_column_facts"
        )
        return self


class DemurrerBuilder(LegalDocumentBuilder):
    """Specialized builder for Demurrer documents"""
    
    def __init__(self):
        super().__init__()
        self.document_type = "Demurrer"
        
    def configure_standard_sections(self):
        """Add standard sections for demurrer"""
        self.add_section("INTRODUCTION", lines_count=10, include_intro_note=True)
        self.add_section("DEMURRER", lines_count=30, include_intro_note=True)
        self.add_section("CONCLUSION", lines_count=5, include_intro_note=True)
        return self