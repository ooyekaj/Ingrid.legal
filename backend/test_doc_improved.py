from docxtpl import DocxTemplate

# Create template (you'll need to create this file with placeholders)
doc = DocxTemplate("template.docx")

# Your detailed context data
context = {
    "court_info": {
        "superior_court": "SUPERIOR COURT OF THE STATE OF CALIFORNIA",
        "county": "COUNTY OF SANTA CLARA",
    },
    "case_info": {
        "plaintiffs": "GILBERT LUNA, et al.",
        "defendant": "GOOGLE LLC",
        "case_no": "24CV434093",
        "related_cases": "Related Google LLC Cases Involving Incognito Mode",
        "assigned_to": "Honorable Charles F. Adams, Department 7",
        "document_title": "PLAINTIFFS' MEMORANDUM IN OPPOSITION TO DEFENDANT GOOGLE LLC'S MOTION FOR SUMMARY JUDGMENT",
        "date_filed": "March 28, 2024",
        "trial_date": "June 24, 2025",
        "motion_date": "June 16, 2025",
        "motion_time": "1:30 p.m.",
        "motion_dept": "7",
    },
    "attorneys_plaintiffs": [
        {
            "firm": "BOIES SCHILLER FLEXNER LLP",
            "attorneys": [
                {
                    "name": "David Boies",
                    "pro_hac_vice": True,
                    "address": "333 Main Street, Armonk, NY 10504",
                    "tel": "(914) 749-8200",
                    "email": "dboies@bsfllp.com",
                    "fax": None
                },
                {
                    "name": "Mark C. Mao",
                    "ca_bar_no": "236165",
                    "address": "44 Montgomery Street, 41st Floor, San Francisco, CA 94104",
                    "tel": "(415) 293-6800",
                    "email": "mmao@bsfllp.com",
                    "fax": "(415) 999-9695"
                },
                {
                    "name": "Beko Reblitz-Richardson",
                    "ca_bar_no": "238027",
                    "address": "44 Montgomery Street, 41st Floor, San Francisco, CA 94104",
                    "tel": "(415) 293-6800",
                    "email": "brichardson@bsfllp.com",
                    "fax": "(415) 999-9695"
                },
                {
                    "name": "Joshua M. Stein",
                    "ca_bar_no": "298856",
                    "address": "44 Montgomery Street, 41st Floor, San Francisco, CA 94104",
                    "tel": "(415) 293-6800",
                    "email": "jstein@bsfllp.com",
                    "fax": "(415) 999-9695"
                },
                {
                    "name": "James Lee",
                    "pro_hac_vice": True,
                    "address": "100 SE 2nd Street, 28th Floor, Miami, FL 33131",
                    "tel": "(305) 539-8400",
                    "email": "jlee@bsfllp.com",
                    "fax": "(305) 539-1307"
                },
                {
                    "name": "Alison L. Anderson",
                    "ca_bar_no": "275334",
                    "address": "2029 Century Park East, Suite 1520, Los Angeles, CA 90067",
                    "tel": "(213) 629-9040",
                    "email": "alanderson@bsfllp.com",
                    "fax": None
                },
                {
                    "name": "M. Logan Wright",
                    "ca_bar_no": "349004",
                    "address": "1401 New York Ave, NW, Washington, D.C. 20005",
                    "tel": "(202) 237-2727",
                    "email": "mwright@bsfllp.com",
                    "fax": "(202) 237-6131"
                },
            ]
        },
        {
            "firm": "MORGAN & MORGAN, P.A.",
            "attorneys": [
                {
                    "name": "John A. Yanchunis",
                    "pro_hac_vice": True,
                    "address": "201 N. Franklin Street, 7th Floor, Tampa, FL 33602",
                    "tel": "(813) 223-5505",
                    "email": "jyanchunis@forthepeople.com",
                    "fax": "(813) 222-4736"
                },
                {
                    "name": "Ryan J. McGee",
                    "pro_hac_vice": True,
                    "address": "201 N. Franklin Street, 7th Floor, Tampa, FL 33602",
                    "tel": "(813) 223-5505",
                    "email": "rmcgee@forthepeople.com",
                    "fax": "(813) 222-4736"
                },
                {
                    "name": "Michael F. Ram",
                    "ca_bar_no": "238027",
                    "address": "201 N. Franklin Street, 7th Floor, Tampa, FL 33602",
                    "tel": "(813) 223-5505",
                    "email": "mram@forthepeople.com",
                    "fax": "(813) 222-4736"
                }
            ]
        }
    ],
    "table_of_contents": [
        {"section": "INTRODUCTION", "page": 8, "level": "main", "number": "I"},
        {"section": "BACKGROUND", "page": 9, "level": "main", "number": "II"},
        {"section": "LEGAL STANDARD", "page": 10, "level": "main", "number": "III"},
        {"section": "ARGUMENT", "page": 10, "level": "main", "number": "IV"},
        {"section": "Plaintiffs' Factual Allegations Are Sufficiently Detailed.", "page": 10, "level": "subsection", "number": "A"},
        {"section": "Plaintiffs' allegations satisfy California's pleading standard.", "page": 10, "level": "subsubsection", "number": "i"},
        {"section": "Google's access to key information also precludes dismissal.", "page": 12, "level": "subsubsection", "number": "ii"},
        {"section": "Plaintiffs Sufficiently Plead Their Penal Code § 631 Claim.", "page": 14, "level": "subsection", "number": "B"},
        {"section": "Plaintiffs Sufficiently Plead Their Penal Code § 632 And Privacy Torts Claims.", "page": 18, "level": "subsection", "number": "C"},
        {"section": "Plaintiffs sufficiently plead their Penal Code § 632 claim.", "page": 18, "level": "subsubsection", "number": "i"},
        {"section": "Plaintiffs sufficiently plead invasion of privacy and intrusion upon seclusion.", "page": 19, "level": "subsubsection", "number": "ii"},
        {"section": "Plaintiffs Allege Facts Establishing UCL & CDAFA Standing.", "page": 20, "level": "subsection", "number": "D"},
        {"section": "Plaintiffs allege facts establishing UCL standing.", "page": 20, "level": "subsubsection", "number": "i"},
        {"section": "Plaintiffs allege facts establishing CDAFA standing.", "page": 22, "level": "subsubsection", "number": "ii"},
        {"section": "Plaintiffs Allege Facts Establishing Viable Contract Damages.", "page": 24, "level": "subsection", "number": "E"},
        {"section": "Plaintiffs Properly Plead Their Unjust Enrichment Claim.", "page": 26, "level": "subsection", "number": "F"},
        {"section": "CONCLUSION", "page": 27, "level": "main", "number": "V"},
    ],
    "sections": {
        "introduction": "INTRODUCTION",
        "background": "BACKGROUND",
        "legal_standard_heading": "LEGAL STANDARD",
        "argument_heading": "ARGUMENT",
        "conclusion_heading": "CONCLUSION",
        "conclusion_text": "For the reasons above, Plaintiffs respectfully request that the Court deny Google's motion for summary judgment in full.",
        "date_signed": "May 27, 2025",
    }
}

# Step 1: Render with docxtpl
doc.render(context)

# Step 2: Save to temporary file
import tempfile
import os
from docx import Document

temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
doc.save(temp_file.name)

# Step 3: Re-open with python-docx to add auto-numbered TOC
final_doc = Document(temp_file.name)

# Find and replace TOC placeholder
toc_inserted = False
for paragraph in final_doc.paragraphs:
    if "TABLE_OF_CONTENTS_PLACEHOLDER" in paragraph.text:
        # Get the paragraph's position to insert TOC entries after it
        parent = paragraph._element.getparent()
        placeholder_index = list(parent).index(paragraph._element)
        
        # Clear the placeholder text
        paragraph.clear()
        
        # Add TOC entries with proper formatting right after the placeholder
        from docx.shared import Inches
        from docx.enum.text import WD_TAB_ALIGNMENT
        
        for i, entry in enumerate(context["table_of_contents"]):
            # Create new paragraph for TOC entry
            new_p = final_doc.add_paragraph()
            
            # Set indentation based on level
            if entry["level"] == "subsection":
                new_p.paragraph_format.left_indent = Inches(0.5)
            elif entry["level"] == "subsubsection":
                new_p.paragraph_format.left_indent = Inches(1.0)
            
            # Add the numbering and section text
            run = new_p.add_run(f"{entry['number']}. {entry['section']}")
            
            # Add tab stops for right-aligned page numbers with dotted leader
            tab_stops = new_p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT, leader=2)  # Dotted leader
            
            # Add tab and page number
            new_p.add_run(f'\t{entry["page"]}')
            
            # Move the paragraph to the correct position
            parent.insert(placeholder_index + 1 + i, new_p._element)
            
            # ADD NEW LINE AFTER EACH SECTION AND SUBSUBSECTION
            if entry["level"] in ["main", "subsubsection"]:
                # Add an empty paragraph for spacing
                empty_p = final_doc.add_paragraph()
                parent.insert(placeholder_index + 2 + i, empty_p._element)
        
        toc_inserted = True
        break

if not toc_inserted:
    print("Warning: TABLE_OF_CONTENTS_PLACEHOLDER not found in template")

# Step 4: Save final document
final_doc.save("luna_v_google_opposition_improved.docx")

# Clean up temporary file
os.unlink(temp_file.name)

print("Document generated as luna_v_google_opposition_improved.docx")
print(f"Case: {context['case_info']['plaintiffs']} v. {context['case_info']['defendant']}")
print(f"Case No.: {context['case_info']['case_no']}")
print(f"Document: {context['case_info']['document_title']}")
print(f"TOC entries: {len(context['table_of_contents'])}")
print("✨ IMPROVEMENT: Added new lines after main sections and subsubsections for better spacing")